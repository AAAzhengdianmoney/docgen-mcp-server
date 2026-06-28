"""
DocGen MCP Server — Generate Word (.docx) and PDF documents via MCP.
Cross-platform: Windows, macOS, and Linux. Multi-font CJK support.

Built with FastMCP. Install: pip install mcp python-docx fpdf2

Workaround for Claude Code CJK bug (#64506):
- sections_file: disk JSON path — the ONLY reliable way for long CJK content.
- sections_b64gz / sections_json_b64: legacy, may still truncate.
"""
from __future__ import annotations

import base64
import gzip
import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from fpdf import FPDF
from fpdf.enums import WrapMode
from mcp.server.fastmcp import FastMCP

mcp = FastMCP(
    name="DocGenerator",
    instructions="Generate Word (.docx) and PDF documents with CJK font support",
)

# ═══════════════════════════════════════════════════════════════════
#  Font Registry — auto-discover all CJK fonts on the system
# ═══════════════════════════════════════════════════════════════════

# display_name → file_path
_FONT_REGISTRY: dict[str, str] = {}

_FONT_CANDIDATES: dict[str, list[str]] = {
    # ── Windows ──
    "微软雅黑":       ["C:/Windows/Fonts/msyh.ttc"],
    "微软雅黑 Bold":  ["C:/Windows/Fonts/msyhbd.ttc"],
    "微软雅黑 Light": ["C:/Windows/Fonts/msyhl.ttc"],
    "宋体":           ["C:/Windows/Fonts/simsun.ttc"],
    "黑体":           ["C:/Windows/Fonts/simhei.ttf"],
    "楷体":           ["C:/Windows/Fonts/simkai.ttf"],
    "仿宋":           ["C:/Windows/Fonts/simfang.ttf"],
    "隶书":           ["C:/Windows/Fonts/SIMLI.TTF"],
    "幼圆":           ["C:/Windows/Fonts/SIMYOU.TTF"],
    "等线":           ["C:/Windows/Fonts/Deng.ttf"],
    "等线 Bold":      ["C:/Windows/Fonts/Dengb.ttf"],
    "等线 Light":     ["C:/Windows/Fonts/Dengl.ttf"],
    "华文仿宋":       ["C:/Windows/Fonts/STFANGSO.TTF"],
    "华文楷体":       ["C:/Windows/Fonts/STKAITI.TTF"],
    "华文宋体":       ["C:/Windows/Fonts/STSONG.TTF"],
    "华文细黑":       ["C:/Windows/Fonts/STXIHEI.TTF"],
    "MingLiU":        ["C:/Windows/Fonts/mingliub.ttc"],
    # ── macOS ──
    "苹方":           ["/System/Library/Fonts/PingFang.ttc"],
    "华文黑体":       ["/System/Library/Fonts/STHeiti Light.ttc"],
    "冬青黑体":       ["/System/Library/Fonts/Hiragino Sans GB.ttc"],
    "Arial Unicode":  ["/Library/Fonts/Arial Unicode.ttf"],
    # ── Linux ──
    "文泉驿正黑":     ["/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"],
    "文泉驿微米黑":   ["/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"],
    "Noto Sans CJK":  [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
    ],
    "Droid Sans":     ["/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"],
    "文鼎上海宋":     ["/usr/share/fonts/truetype/arphic/uming.ttc"],
    "WenQuanYi Zen Hei":    ["/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"],
    "WenQuanYi Micro Hei":  ["/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"],
}

for name, paths in _FONT_CANDIDATES.items():
    for p in paths:
        if Path(p).exists():
            _FONT_REGISTRY[name] = p
            break

# Short aliases for quick access in sections
_FONT_ALIASES: dict[str, str] = {
    "yahei":     "微软雅黑",
    "雅黑":      "微软雅黑",
    "song":      "宋体",
    "simsun":    "宋体",
    "hei":       "黑体",
    "simhei":    "黑体",
    "kai":       "楷体",
    "simkai":    "楷体",
    "fang":      "仿宋",
    "simfang":   "仿宋",
    "li":        "隶书",
    "simli":     "隶书",
    "you":       "幼圆",
    "simyou":    "幼圆",
    "deng":      "等线",
    "dengbold":  "等线 Bold",
    "denglight": "等线 Light",
    "stfangsong": "华文仿宋",
    "stkaiti":   "华文楷体",
    "stsong":    "华文宋体",
    "stxihei":   "华文细黑",
    "pingfang":  "苹方",
    "noto":      "Noto Sans CJK",
}

_ANY_CJK = bool(_FONT_REGISTRY)

# ── Default font selection ──

def _pick_default_body() -> tuple[str, str]:
    for name in ["微软雅黑", "苹方", "文泉驿微米黑", "Noto Sans CJK"]:
        if name in _FONT_REGISTRY:
            return name, _FONT_REGISTRY[name]
    for name, path in _FONT_REGISTRY.items():
        return name, path
    return ("Arial", "")

def _pick_default_heading() -> tuple[str, str]:
    for name in ["黑体", "苹方", "文泉驿正黑", "Noto Sans CJK"]:
        if name in _FONT_REGISTRY:
            return name, _FONT_REGISTRY[name]
    return _pick_default_body()

_DEFAULT_BODY_NAME, _DEFAULT_BODY_PATH = _pick_default_body()
_DEFAULT_HEADING_NAME, _DEFAULT_HEADING_PATH = _pick_default_heading()

# Legacy compat
_CJK_FONT = _DEFAULT_BODY_PATH or None


def _resolve_font(name: str | None) -> str | None:
    """Resolve font name (or alias) → display name in registry. Returns None if not found."""
    if not name:
        return None
    if name in _FONT_REGISTRY:
        return name
    # Try aliases (case-insensitive)
    key = name.lower()
    if key in _FONT_ALIASES:
        resolved = _FONT_ALIASES[key]
        if resolved in _FONT_REGISTRY:
            return resolved
    # Try case-insensitive match against registry keys
    for reg_name in _FONT_REGISTRY:
        if reg_name.lower() == key:
            return reg_name
    return None


def _resolve_font_path(name: str | None) -> str:
    """Resolve font name → file path. Returns default body path if not found."""
    resolved = _resolve_font(name)
    if resolved:
        return _FONT_REGISTRY[resolved]
    return _DEFAULT_BODY_PATH


# ═══════════════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════════════

def _decode_sections(
    sections: list[dict] = [],
    sections_json_b64: str = "",
    sections_b64gz: str = "",
    sections_file: str = "",
) -> list[dict]:
    """Decode sections from direct list, base64 JSON, or gzip+base64 JSON."""
    if sections_file:
        return json.loads(Path(sections_file).read_text(encoding="utf-8"))
    if sections_b64gz:
        try:
            raw = base64.b64decode(sections_b64gz)
            return json.loads(gzip.decompress(raw).decode("utf-8"))
        except Exception:
            pass
    if sections_json_b64:
        try:
            decoded = base64.b64decode(sections_json_b64).decode("utf-8")
            return json.loads(decoded)
        except Exception:
            pass
    return sections or []


def _set_docx_font(style, font_name: str):
    """Set both western and east-asia font on a paragraph style via XML."""
    from docx.oxml.ns import qn
    from lxml import etree

    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _set_run_font(run, font_name: str):
    """Set CJK font on a single run via XML (for headings/title)."""
    from docx.oxml.ns import qn
    from lxml import etree

    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _make_docx_body_paragraph(doc, body: str, font_name: str):
    """Add body paragraphs with a specific font. Multi-line text → multiple paragraphs."""
    for line in body.strip().split("\n"):
        text = line.strip()
        if text:
            p = doc.add_paragraph(text)
            for run in p.runs:
                _set_run_font(run, font_name)


def _make_docx_bullet(doc, body: str, font_name: str):
    """Add bullet items with a specific font."""
    for line in body.strip().split("\n"):
        if line.strip():
            p = doc.add_paragraph(line.strip(), style="List Bullet")
            for run in p.runs:
                _set_run_font(run, font_name)


def _make_docx_numbered(doc, body: str, font_name: str):
    """Add numbered items with a specific font."""
    for line in body.strip().split("\n"):
        if line.strip():
            p = doc.add_paragraph(line.strip(), style="List Number")
            for run in p.runs:
                _set_run_font(run, font_name)


# ── Rich text (markdown-like inline formatting) ──

import re

_MD_PAT = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")

def _parse_markdown_spans(text: str) -> list[dict]:
    """Parse **bold** and *italic* spans. Returns [{"text": ..., "bold": bool, "italic": bool}]."""
    spans = []
    pos = 0
    for m in _MD_PAT.finditer(text):
        if m.start() > pos:
            spans.append({"text": text[pos:m.start()], "bold": False, "italic": False})
        if m.group(2):  # **...**
            spans.append({"text": m.group(2), "bold": True, "italic": False})
        else:           # *...*
            spans.append({"text": m.group(3), "bold": False, "italic": True})
        pos = m.end()
    if pos < len(text):
        spans.append({"text": text[pos:], "bold": False, "italic": False})
    return spans


def _make_docx_rich_paragraph(doc, body: str, font_name: str):
    """Add body paragraphs with **bold** / *italic* inline formatting. Multi-line."""
    for line in body.strip().split("\n"):
        text = line.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        for span in _parse_markdown_spans(text):
            run = p.add_run(span["text"])
            run.bold = span["bold"]
            if span["italic"]:
                run.underline = True  # CJK fonts have no italic glyphs; underline instead
            _set_run_font(run, font_name)


# ── Table ──

def _make_docx_table(doc, table_def: dict, font_name: str):
    """Add a styled table with CJK font support."""
    headers = table_def.get("headers", [])
    rows = table_def.get("rows", [])
    if not headers and not rows:
        return

    ncols = len(headers) or (len(rows[0]) if rows else 1)
    table = doc.add_table(rows=1 + len(rows), cols=ncols, style="Table Grid")

    def _fill_cell(cell, text: str):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        _set_run_font(run, font_name)

    # Header row
    for i, h in enumerate(headers):
        _fill_cell(table.rows[0].cells[i], h)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _fill_cell(table.rows[ri + 1].cells[ci], val)

    # Add a blank paragraph after the table for spacing
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════
#  PDF helpers
# ═══════════════════════════════════════════════════════════════════

# Track which fonts have been registered per PDF instance
_pdf_registered_fonts: dict[int, dict[str, str]] = {}  # id(pdf) → {display_name: fpdf2_key}


def _ensure_pdf_font(pdf: FPDF, display_name: str) -> str:
    """Register a CJK font in the PDF if not already there. Returns fpdf2 font key."""
    pdf_id = id(pdf)
    if pdf_id not in _pdf_registered_fonts:
        _pdf_registered_fonts[pdf_id] = {}
    registered = _pdf_registered_fonts[pdf_id]

    if display_name in registered:
        return registered[display_name]

    path = _resolve_font_path(display_name)
    if not path:
        # No CJK font at all — use Helvetica
        return "Helvetica"

    # Use a short key for fpdf2 (ASCII-safe)
    key = f"f{len(registered)}"
    pdf.add_font(key, "", path)
    registered[display_name] = key
    return key


def _pdf_write_text(pdf: FPDF, text: str, font_display: str, size: int,
                    align: str = "L", new_x: str = "LMARGIN", new_y: str = "NEXT"):
    """Write text with a specific CJK font in PDF."""
    if _ANY_CJK:
        key = _ensure_pdf_font(pdf, font_display)
        pdf.set_font(key, "", size)
    else:
        pdf.set_font("Helvetica", "", size)

    if align == "C":
        pdf.cell(0, size * 0.6, text, align="C", new_x=new_x, new_y=new_y)
    else:
        bw = pdf.w - pdf.l_margin - pdf.r_margin
        lines = pdf.multi_cell(
            bw, size * 0.55, text,
            dry_run=True, output="LINES", wrapmode=WrapMode.CHAR,
        )
        for ln in lines:
            pdf.cell(bw, size * 0.55, ln, new_x="LMARGIN", new_y="NEXT")


def _pdf_write_table(pdf: FPDF, table_def: dict, font_display: str):
    """Draw a table with CJK font in PDF."""
    headers = table_def.get("headers", [])
    rows = table_def.get("rows", [])
    if not headers and not rows:
        return

    ncols = len(headers) or (len(rows[0]) if rows else 1)
    col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / ncols
    line_h = 7
    font_size = 9

    if _ANY_CJK:
        key = _ensure_pdf_font(pdf, font_display)

    def _cell(text: str, w: float, bold: bool = False):
        if _ANY_CJK:
            pdf.set_font(key, "", font_size)
        else:
            pdf.set_font("Helvetica", "B" if bold else "", font_size)
        pdf.cell(w, line_h, str(text), border=1, align="C")

    # Header row
    for h in headers:
        _cell(h, col_w, bold=True)
    pdf.ln(line_h)

    # Data rows
    for row in rows:
        for ci, val in enumerate(row):
            _cell(val, col_w)
        pdf.ln(line_h)
    pdf.ln(4)


class _PDF(FPDF):
    def __init__(self, header_text: str = "", footer_text: str = ""):
        super().__init__()
        self._header_text = header_text
        self._footer_text = footer_text
        # Register a CJK font for header/footer text
        self._hf_key = None
        if _ANY_CJK and _DEFAULT_BODY_PATH:
            self._hf_key = "hf0"
            self.add_font(self._hf_key, "", _DEFAULT_BODY_PATH)

    def header(self):
        if not self._header_text:
            return
        if self._hf_key:
            self.set_font(self._hf_key, "", 8)
        else:
            self.set_font("Helvetica", "", 8)
        self.cell(0, 10, self._header_text, align="C", new_x="LMARGIN", new_y="NEXT")

    def footer(self):
        if self._hf_key:
            self.set_font(self._hf_key, "", 8)
        else:
            self.set_font("Helvetica", "", 8)
        self.set_y(-15)
        if self._footer_text:
            self.cell(0, 10, self._footer_text + f"  |  Page {self.page_no()}/{{nb}}", align="C")
        else:
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


# ═══════════════════════════════════════════════════════════════════
#  Tools
# ═══════════════════════════════════════════════════════════════════

@mcp.tool()
def list_fonts() -> str:
    """List all available CJK fonts discovered on this system.
    Use this to see what font names you can pass to generate_docx / generate_pdf
    via the 'font' and 'heading_font' fields in each section dict.
    """
    if not _FONT_REGISTRY:
        return "No CJK fonts found. Install fonts or use ASCII-only content."

    lines = ["Available CJK fonts:", ""]
    # Group by platform
    win = {k: v for k, v in _FONT_REGISTRY.items() if "Windows" in v}
    mac = {k: v for k, v in _FONT_REGISTRY.items() if "System/Library" in v or "Library/Fonts" in v}
    linux = {k: v for k, v in _FONT_REGISTRY.items() if "/usr/" in v}

    def _fmt(section: str, fonts: dict):
        if not fonts:
            return
        lines.append(f"  {section}:")
        for name, path in sorted(fonts.items()):
            lines.append(f"    • {name}")
        lines.append("")

    _fmt("🪟 Windows", win)
    _fmt("🍎 macOS", mac)
    _fmt("🐧 Linux", linux)

    lines.append(f"  Default body font: {_DEFAULT_BODY_NAME}")
    lines.append(f"  Default heading font: {_DEFAULT_HEADING_NAME}")
    lines.append("")
    lines.append("Aliases: " + ", ".join(sorted(_FONT_ALIASES.keys())))
    return "\n".join(lines)


@mcp.tool()
def generate_docx(
    output_path: str,
    title: str = "",
    sections: list[dict] = [],
    title_font: str = "",
    heading_font: str = "",
    header_text: str = "",
    footer_text: str = "",
    sections_json_b64: str = "",
    sections_b64gz: str = "",
    sections_file: str = "",
) -> str:
    """Generate a Word (.docx) document with title and sections.

    Parameters:
        output_path: Absolute path to save the .docx file.
        title: Document title (centered, large heading).
        title_font: Font for the main title (default: Hei/SimHei).
        heading_font: Default font for all section headings (default: Hei/SimHei).
            Individual sections can override with their own heading_font.
        header_text: Optional page header text.
        footer_text: Optional page footer text (appended before page number).
        sections: List of section dicts (ASCII-safe only, for short content).
        sections_file: Path to a JSON file on disk containing the sections array.
                      BEST for CJK/large content — avoids MCP argument serialization bugs.
        sections_json_b64: Base64-encoded JSON of sections array (legacy).
        sections_b64gz: gzip-compressed + base64-encoded sections JSON (legacy).
        sections[].heading: Section heading.
        sections[].body: Section body text. Supports **bold** and *italic* inline.
        sections[].style: "normal" (default), "bullet", "numbered", "table", "divider", "image".
        sections[].level: Heading level 1/2/3 (default: 1).
        sections[].font: Font for body text in this section (default: 微软雅黑/MS YaHei).
            Use list_fonts() to see available fonts. Aliases like "kai", "fang", "hei" work.
        sections[].heading_font: Font for this section's heading (default: 黑体/SimHei).
        sections[].table: {"headers": [...], "rows": [[...], ...]} for table style.
        sections[].image: Path to an image file (local) for image style.
        sections[].image_width: Image width in mm (optional, for image style).
    """
    sections = _decode_sections(sections, sections_json_b64, sections_b64gz, sections_file)

    # Resolve global heading font
    global_heading = _resolve_font(heading_font) or _DEFAULT_HEADING_NAME
    tfont = _resolve_font(title_font) or global_heading

    doc = Document()

    # ── Normal style defaults ──
    doc.styles["Normal"].font.size = Pt(11)
    _set_docx_font(doc.styles["Normal"], _DEFAULT_BODY_NAME)

    # ── Header / Footer ──
    if header_text:
        h = doc.sections[0].header
        h.paragraphs[0].clear()
        run = h.paragraphs[0].add_run(header_text)
        run.font.size = Pt(8)
        h.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_text:
        f = doc.sections[0].footer
        f.paragraphs[0].clear()
        run = f.paragraphs[0].add_run(footer_text)
        run.font.size = Pt(8)
        f.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Title ──
    if title:
        p = doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            _set_run_font(run, tfont)

    # ── Sections ──
    for s in sections:
        heading = s.get("heading", "")
        body = s.get("body", "")
        sn = s.get("style", "normal")
        level = s.get("level", 1)

        # Per-section font overrides
        body_font = _resolve_font(s.get("font")) or _DEFAULT_BODY_NAME
        sec_heading_font = _resolve_font(s.get("heading_font")) or global_heading
        table_font = _resolve_font(s.get("table_font")) or body_font

        # ── Divider ──
        if sn == "divider":
            p = doc.add_paragraph("─" * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            continue

        # ── Image ──
        if sn == "image":
            img_path = s.get("image", "")
            if img_path and Path(img_path).exists():
                w = s.get("image_width")
                kw = {"width": Inches(w / 25.4)} if w else {}
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, **kw)
            continue

        # ── Heading (with level) ──
        if heading:
            h = doc.add_heading(heading, level=level)
            for run in h.runs:
                _set_run_font(run, sec_heading_font)

        if not body and sn != "table":
            continue

        # ── Table ──
        if sn == "table":
            td = s.get("table", {})
            if td:
                _make_docx_table(doc, td, table_font)
            continue

        # ── Body text styles ──
        if sn == "bullet":
            _make_docx_bullet(doc, body, body_font)
        elif sn == "numbered":
            _make_docx_numbered(doc, body, body_font)
        else:
            _make_docx_rich_paragraph(doc, body, body_font)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return f"OK: {out.resolve()} ({out.stat().st_size} bytes)"


@mcp.tool()
def generate_pdf(
    output_path: str,
    title: str = "",
    sections: list[dict] = [],
    title_font: str = "",
    heading_font: str = "",
    header_text: str = "",
    footer_text: str = "",
    sections_json_b64: str = "",
    sections_b64gz: str = "",
    sections_file: str = "",
) -> str:
    """Generate a PDF document with title and sections.

    Parameters:
        output_path: Absolute path to save the .pdf file.
        title: Document title.
        title_font: Font for the main title (default: Hei/SimHei).
        heading_font: Default font for all section headings.
        header_text: Optional page header text.
        footer_text: Optional page footer text (appended before page number).
        sections: List of section dicts (ASCII-safe only, for short content).
        sections_file: Path to a JSON file on disk containing the sections array.
                      BEST for CJK/large content — avoids MCP argument serialization bugs.
        sections_json_b64: Base64-encoded JSON of sections array (legacy).
        sections_b64gz: gzip-compressed + base64-encoded sections JSON (legacy).
        sections[].heading: Section heading.
        sections[].body: Section body text. **bold** / *italic* markers are stripped in PDF.
        sections[].font: Font for body text in this section.
            Use list_fonts() to see available fonts. Aliases like "kai", "fang", "hei" work.
        sections[].heading_font: Font for this section's heading.
        sections[].style: "normal" (default), "table", "divider", "image", "bullet", "numbered".
        sections[].level: Heading level 1/2/3 (default: 1).
        sections[].table: {"headers": [...], "rows": [[...], ...]} for table style.
        sections[].image: Path to an image file (local) for image style.
        sections[].image_width: Image width in mm (optional).
    """
    sections = _decode_sections(sections, sections_json_b64, sections_b64gz, sections_file)

    global_heading = _resolve_font(heading_font) or _DEFAULT_HEADING_NAME
    tfont = _resolve_font(title_font) or global_heading
    default_body = _DEFAULT_BODY_NAME

    pdf = _PDF(header_text=header_text, footer_text=footer_text)
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    if title:
        _pdf_write_text(pdf, title, tfont, 18, align="C")
        pdf.ln(5)

    for s in sections:
        heading = s.get("heading", "")
        body = s.get("body", "")
        sn = s.get("style", "normal")
        level = s.get("level", 1)

        body_font = _resolve_font(s.get("font")) or default_body
        sec_heading_font = _resolve_font(s.get("heading_font")) or global_heading
        table_font = _resolve_font(s.get("table_font")) or body_font

        # ── Divider ──
        if sn == "divider":
            y = pdf.get_y()
            bw = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.line(pdf.l_margin, y, pdf.l_margin + bw, y)
            pdf.ln(4)
            continue

        # ── Image ──
        if sn == "image":
            img_path = s.get("image", "")
            if img_path and Path(img_path).exists():
                w = s.get("image_width")
                max_w = pdf.w - pdf.l_margin - pdf.r_margin
                if w is None:
                    w = max_w
                x = pdf.l_margin + (max_w - w) / 2
                pdf.image(img_path, x=x, w=w)
                pdf.ln(4)
            continue

        # ── Heading with level ──
        if heading:
            sizes = {1: 14, 2: 12, 3: 11}
            _pdf_write_text(pdf, heading, sec_heading_font, sizes.get(level, 14))
            pdf.ln(2)

        if not body and sn != "table":
            continue

        # ── Table ──
        if sn == "table":
            td = s.get("table", {})
            if td:
                _pdf_write_table(pdf, td, table_font)
            continue

        # ── Body: strip markdown markers for PDF (no bold/italic in CJK fonts) ──
        clean_body = _MD_PAT.sub(lambda m: m.group(2) or m.group(3), body)

        if sn == "bullet":
            for line in clean_body.strip().split("\n"):
                if line.strip():
                    _pdf_write_text(pdf, "  • " + line.strip(), body_font, 11)
        elif sn == "numbered":
            for i, line in enumerate(clean_body.strip().split("\n"), 1):
                if line.strip():
                    _pdf_write_text(pdf, f"  {i}. {line.strip()}", body_font, 11)
        else:
            _pdf_write_text(pdf, clean_body, body_font, 11)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        tmppath = Path(f.name)
    pdf.output(str(tmppath))
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(tmppath.read_bytes())
    return f"OK: {out.resolve()} ({out.stat().st_size} bytes)"


if __name__ == "__main__":
    mcp.run(transport="stdio")
